import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from docx import Document
from tqdm import tqdm
import time
import re
import argparse

# === DICTIONARY NAMA BAHASA KE KODE ===
language_map = {
    "indonesia": "id", "bahasa indonesia": "id",
    "english": "en", "inggris": "en",
    "france": "fr", "french": "fr", "prancis": "fr",
    "spanish": "es", "spanyol": "es",
    "japanese": "ja", "jepang": "ja",
    "german": "de", "jerman": "de",
    "arabic": "ar", "arab": "ar",
    "korean": "ko", "korea": "ko",
    "chinese": "zh-cn", "tiongkok": "zh-cn", "mandarin": "zh-cn"
}

def get_lang_code(user_input):
    """Konversi nama bahasa ke kode ISO."""
    key = user_input.strip().lower()
    return language_map.get(key, "id")  # default ke Indonesia

# === PARSER ARGUMEN ===
parser = argparse.ArgumentParser(description="Terjemahkan file PDF ke DOCX.")
parser.add_argument("--input", help="Nama file PDF sumber", required=False)
parser.add_argument("--output", help="Nama file DOCX hasil", required=False)
parser.add_argument("--lang", help="Bahasa tujuan (contoh: Indonesia, Inggris, Prancis)", required=False)
args = parser.parse_args()

# === INPUT INTERAKTIF JIKA TIDAK ADA ARGUMEN ===
if not args.input:
    args.input = input("📂 Masukkan nama file PDF (mis. file.pdf): ").strip()
if not args.input.lower().endswith(".pdf"):
    args.input += ".pdf"

if not args.output:
    args.output = input("💾 Nama file output DOCX (mis. hasil_terjemahan.docx): ").strip()
if not args.output.lower().endswith(".docx"):
    args.output += ".docx"

if not args.lang:
    args.lang = input("🌍 Masukkan bahasa tujuan (contoh: Indonesia, Inggris, Jepang): ").strip()

target_lang = get_lang_code(args.lang)

print(f"\n🚀 Menerjemahkan '{args.input}' ke bahasa '{args.lang.capitalize()}' ({target_lang}) ...\n")

# === KONFIGURASI PENERJEMAH ===
translator = GoogleTranslator(source='auto', target=target_lang)
doc = fitz.open(args.input)
word_doc = Document()

print(f"📘 Jumlah halaman: {len(doc)}\n")

def bersihkan_teks(teks):
    """Hilangkan karakter aneh biar gak bikin error."""
    teks = re.sub(r'[^\x00-\x7F]+', ' ', teks)
    teks = re.sub(r'\s+', ' ', teks)
    return teks.strip()

def safe_translate(text):
    """Terjemahkan teks per paragraf dengan retry otomatis."""
    text = bersihkan_teks(text)
    if not text or text.startswith("http"):
        return text

    hasil = ""
    paragraphs = text.split("\n")
    for paragraf in paragraphs:
        paragraf = paragraf.strip()
        if not paragraf:
            continue

        for attempt in range(3):
            try:
                translated = translator.translate(paragraf)
                hasil += translated + "\n"
                break
            except Exception as e:
                print(f"⚠️ Gagal terjemah (percobaan {attempt+1}/3): {e}")
                time.sleep(3)
        else:
            hasil += paragraf + "\n"

    return hasil

# === PROSES TERJEMAHAN PER HALAMAN ===
for i, page in enumerate(tqdm(doc, desc="📄 Menerjemahkan halaman")):
    text = page.get_text("text")
    if not text.strip():
        continue

    translated = safe_translate(text)
    word_doc.add_heading(f"Halaman {i+1}", level=2)
    word_doc.add_paragraph(translated)

# === SIMPAN HASIL ===
word_doc.save(args.output)
print(f"\n✅ Terjemahan selesai! File disimpan sebagai: {args.output}")
